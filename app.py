"""
Mobile Resume Generator - Containerized app for generating resumes from job descriptions
Just paste a JD and get: Resume PDF + Response Email
"""

import os
import sys
import json
import re
import requests
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS

# Add current directory to path for imports
sys.path.insert(0, str(Path(__file__).parent))

app = Flask(__name__)
CORS(app)

# File-based storage for multi-worker compatibility
STORAGE_FILE = Path(__file__).parent / "output" / "last_session.json"

def save_session(data):
    """Save session data to file"""
    STORAGE_FILE.parent.mkdir(exist_ok=True)
    with open(STORAGE_FILE, 'w') as f:
        json.dump(data, f)

def load_session():
    """Load session data from file"""
    if STORAGE_FILE.exists():
        with open(STORAGE_FILE, 'r') as f:
            return json.load(f)
    return {}

# User Profile (your details)
USER_PROFILE = {
    "name": "Muhammad Kasim Naina Mohammed",
    "email": "mohammedmazhermd@gmail.com",
    "phone": "(510) 771-4493",
    "linkedin": "https://www.linkedin.com/in/muhammad-kasim-0b297416/",
    "location": "San Jose, CA"
}

# LLM Configuration - Cerebras API key (fallback if env not set)
DEFAULT_CEREBRAS_KEY = "csk-33tdk3dd6ktyyjecw65p4x5wdx4exdkrrh5d3jndfcj6cdxp"

LLM_PROVIDERS = {
    "cerebras": {
        "url": "https://api.cerebras.ai/v1/chat/completions",
        "model": "llama-3.3-70b",
        "env_key": "CEREBRAS_API_KEY",
        "default_key": DEFAULT_CEREBRAS_KEY
    },
    "groq": {
        "url": "https://api.groq.com/openai/v1/chat/completions",
        "model": "llama-3.3-70b-versatile",
        "env_key": "GROQ_API_KEY"
    },
    "openrouter": {
        "url": "https://openrouter.ai/api/v1/chat/completions",
        "model": "meta-llama/llama-3.3-70b-instruct:free",
        "env_key": "OPENROUTER_API_KEY"
    }
}

def get_llm_config():
    """Get first available LLM provider"""
    for provider, config in LLM_PROVIDERS.items():
        api_key = os.environ.get(config["env_key"], "") or config.get("default_key", "")
        if api_key:
            return {"provider": provider, **config, "api_key": api_key}
    raise ValueError("No LLM API key configured")

def call_llm(prompt, max_tokens=4096):
    """Call LLM API"""
    config = get_llm_config()
    headers = {"Authorization": f"Bearer {config['api_key']}", "Content-Type": "application/json"}
    payload = {
        "model": config["model"],
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7,
        "max_tokens": max_tokens
    }
    response = requests.post(config["url"], headers=headers, json=payload, timeout=120)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]

def extract_job_info(job_description):
    """Extract job title, company, location from JD using LLM"""
    prompt = f"""Extract the following from this job description and return as JSON:
- job_title: The job title/position
- company: The company name (or recruiter company if client not mentioned)
- location: The job location
- job_type: remote/hybrid/onsite
- key_skills: List of top 5-7 required skills

Job Description:
{job_description}

Return ONLY valid JSON, no other text:
{{"job_title": "...", "company": "...", "location": "...", "job_type": "...", "key_skills": [...]}}"""
    
    result = call_llm(prompt, max_tokens=500)
    # Extract JSON from response
    try:
        match = re.search(r'\{.*\}', result, re.DOTALL)
        if match:
            return json.loads(match.group())
    except:
        pass
    return {"job_title": "Software Engineer", "company": "Company", "location": "Remote", "job_type": "hybrid", "key_skills": []}

def load_base_resume():
    """Load base resume content"""
    resume_path = Path(__file__).parent / "data" / "resume.txt"
    if resume_path.exists():
        return resume_path.read_text()
    # Fallback to parent project resume
    alt_path = Path(__file__).parent.parent / "job-application-mcp-server" / "resumes" / "default_resume.md"
    if alt_path.exists():
        return alt_path.read_text()
    return ""

@app.route('/')
def index():
    return render_template('index.html', profile=USER_PROFILE)

@app.route('/api/generate', methods=['POST'])
def generate():
    """Main endpoint - takes JD, returns resume + email"""
    data = request.json
    job_description = data.get('job_description', '')

    if not job_description:
        return jsonify({"error": "Job description required"}), 400

    try:
        # Step 1: Extract job info
        job_info = extract_job_info(job_description)

        # Step 2: Generate tailored resume
        resume_data = generate_resume_data(job_info, job_description)

        # Store for download (file-based for multi-worker)
        save_session({
            'resume': resume_data,
            'job_info': job_info,
            'jd': job_description
        })

        # Step 3: Generate response email
        email = generate_response_email(job_info, job_description)

        return jsonify({
            "success": True,
            "job_info": job_info,
            "email": email,
            "resume_url": "/api/resume/download"
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

def generate_resume_data(job_info, job_description):
    """Generate structured resume data using the existing generate_resume module"""
    try:
        # Import from local generate_resume.py
        import generate_resume as gen
        import importlib
        importlib.reload(gen)

        # Load resume with current role config
        resume = gen.load_resume()

        job_title = job_info.get('job_title', 'Software Engineer')
        company = job_info.get('company', 'Company')

        print(f"🤖 Generating resume for: {job_title} at {company}")

        # Use the existing get_structured_resume function
        data = gen.get_structured_resume(job_title, company, job_description, resume)

        if data:
            # Override contact info with user profile
            data['name'] = USER_PROFILE['name']
            data['email'] = USER_PROFILE['email']
            data['phone'] = USER_PROFILE['phone']
            data['linkedin'] = USER_PROFILE['linkedin']
            data['location'] = USER_PROFILE['location']
            return data

        return None
    except Exception as e:
        print(f"❌ Resume generation error: {e}")
        import traceback
        traceback.print_exc()
        return None

def generate_response_email(job_info, job_description):
    """Generate professional response email"""
    prompt = f"""Write a professional, concise job application response email.

Job Title: {job_info.get('job_title')}
Company: {job_info.get('company')}
Location: {job_info.get('location')}
Key Skills Required: {', '.join(job_info.get('key_skills', []))}

Candidate Info:
- Name: {USER_PROFILE['name']}
- 12+ years Java/microservices experience
- Currently at Intuit, previously at Amazon, PayPal, Citi
- Expert in: Java, Spring Boot, Microservices, Kubernetes, Kafka, AWS

Write a SHORT (under 150 words) professional email expressing interest. Include:
1. Thank them for reaching out
2. Brief highlight of relevant experience (2-3 bullet points max)
3. Mention availability for interview
4. Professional signature with contact info

Return ONLY the email body, no subject line."""

    return call_llm(prompt, max_tokens=500)

@app.route('/api/resume/download')
def download_resume():
    """Download generated resume as PDF or DOCX"""
    format_type = request.args.get('format', 'pdf')
    company = request.args.get('company', 'resume')

    # Load from file-based session
    session = load_session()
    resume_data = session.get('resume')
    if not resume_data:
        return jsonify({"error": "No resume generated yet. Please generate first!"}), 400

    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)

    filename = f"resume_{company}_{format_type}"
    filepath = output_dir / f"{filename}.{format_type}"

    if format_type == 'pdf':
        generate_pdf(resume_data, str(filepath))
    else:
        generate_docx(resume_data, str(filepath))

    return send_file(filepath, as_attachment=True, download_name=f"resume_{company}.{format_type}")

def generate_pdf(data, output_path):
    """Generate PDF resume"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    doc = SimpleDocTemplate(output_path, pagesize=letter, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    story = []

    # Name
    name_style = ParagraphStyle('Name', parent=styles['Title'], fontSize=18, spaceAfter=6)
    story.append(Paragraph(data.get('name', ''), name_style))

    # Contact
    contact = f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('location', '')}"
    story.append(Paragraph(contact, styles['Normal']))
    story.append(Spacer(1, 12))

    # Summary
    if data.get('summary'):
        story.append(Paragraph("<b>PROFESSIONAL SUMMARY</b>", styles['Heading2']))
        story.append(Paragraph(data['summary'], styles['Normal']))
        story.append(Spacer(1, 12))

    # Skills
    if data.get('skills'):
        story.append(Paragraph("<b>TECHNICAL SKILLS</b>", styles['Heading2']))
        for category, skills in data['skills'].items():
            story.append(Paragraph(f"<b>{category}:</b> {', '.join(skills)}", styles['Normal']))
        story.append(Spacer(1, 12))

    # Experience
    story.append(Paragraph("<b>PROFESSIONAL EXPERIENCE</b>", styles['Heading2']))
    for job in data.get('experience', []):
        story.append(Paragraph(f"<b>{job.get('title', '')}</b> | {job.get('company', '')}", styles['Normal']))
        story.append(Paragraph(f"{job.get('dates', '')} | {job.get('location', '')}", styles['Normal']))
        for bullet in job.get('bullets', [])[:4]:
            story.append(Paragraph(f"• {bullet}", styles['Normal']))
        story.append(Spacer(1, 8))

    doc.build(story)
    print(f"✅ PDF saved: {output_path}")

def generate_docx(data, output_path):
    """Generate DOCX resume"""
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()

    # Name
    name = doc.add_heading(data.get('name', ''), 0)
    name.alignment = 1

    # Contact
    contact = doc.add_paragraph()
    contact.add_run(f"{data.get('email', '')} | {data.get('phone', '')} | {data.get('location', '')}")
    contact.alignment = 1

    # Summary
    if data.get('summary'):
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(data['summary'])

    # Skills
    if data.get('skills'):
        doc.add_heading('Technical Skills', level=1)
        for category, skills in data['skills'].items():
            p = doc.add_paragraph()
            p.add_run(f"{category}: ").bold = True
            p.add_run(', '.join(skills))

    # Experience
    doc.add_heading('Professional Experience', level=1)
    for job in data.get('experience', []):
        p = doc.add_paragraph()
        p.add_run(f"{job.get('title', '')}").bold = True
        p.add_run(f" | {job.get('company', '')}")
        doc.add_paragraph(f"{job.get('dates', '')} | {job.get('location', '')}")
        for bullet in job.get('bullets', [])[:4]:
            doc.add_paragraph(f"• {bullet}")

    doc.save(output_path)
    print(f"✅ DOCX saved: {output_path}")

if __name__ == '__main__':
    print("\n📱 Mobile Resume Generator")
    print("   http://localhost:8080\n")
    app.run(host='0.0.0.0', port=8080, debug=True)

