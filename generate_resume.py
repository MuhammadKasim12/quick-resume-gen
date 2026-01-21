#!/usr/bin/env python3
"""
Professional Resume Generator with AI Customization
Generates PDF and DOCX resumes using industry-standard templates

Supports multiple LLM providers:
- groq: Groq API (default) - llama-3.3-70b-versatile
- cerebras: Cerebras API - llama-3.3-70b (fast inference)
- openrouter: OpenRouter API - various free models

Set LLM_PROVIDER environment variable to switch providers.
Set corresponding API key: GROQ_API_KEY, CEREBRAS_API_KEY, OPENROUTER_API_KEY
"""

import requests
import json
import sys
import os
from pathlib import Path

# Add parent directory for resume handler
sys.path.insert(0, str(Path(__file__).parent.parent / 'job-application-mcp-server'))

# LLM Provider Configuration
LLM_PROVIDERS = {
    "cerebras": {
        "name": "Cerebras",
        "url": "https://api.cerebras.ai/v1/chat/completions",
        "model": "llama-3.3-70b",
        "env_key": "CEREBRAS_API_KEY",
        "default_key": ""
    },
    "groq": {
        "name": "Groq",
        "url": "https://api.groq.com/openai/v1/chat/completions",
        "model": "llama-3.3-70b-versatile",
        "env_key": "GROQ_API_KEY",
        "default_key": ""
    },
    "openrouter": {
        "name": "OpenRouter",
        "url": "https://openrouter.ai/api/v1/chat/completions",
        "model": "meta-llama/llama-3.3-70b-instruct:free",
        "env_key": "OPENROUTER_API_KEY",
        "default_key": ""
    }
}

def get_llm_config():
    """Get the current LLM provider configuration - prioritize Cerebras"""
    # Try providers in order: cerebras, groq, openrouter
    for provider_name in ["cerebras", "groq", "openrouter"]:
        config = LLM_PROVIDERS[provider_name]
        api_key = os.environ.get(config["env_key"], config.get("default_key", ""))
        if api_key:
            print(f"✅ Using {config['name']} LLM provider")
            return {
                "provider": provider_name,
                "name": config["name"],
                "url": config["url"],
                "model": config["model"],
                "api_key": api_key
            }

    raise ValueError("❌ No API key found. Set CEREBRAS_API_KEY, GROQ_API_KEY, or OPENROUTER_API_KEY")

def load_resume():
    """Load resume from local resumes folder"""
    try:
        from resume_handler import ResumeHandler
        # Try local resumes folder first (for Render/Docker deployment)
        local_resumes = Path(__file__).parent / 'resumes'
        if local_resumes.exists():
            handler = ResumeHandler(str(local_resumes))
        else:
            # Fallback to job-application-mcp-server path (for local dev)
            handler = ResumeHandler(str(Path(__file__).parent.parent / 'job-application-mcp-server' / 'resumes'))

        resume = handler.get_default_resume()
        resume_content = resume.content if resume else ""
        print(f"📄 Loaded resume: {len(resume_content)} chars")

        # Load current role override if exists
        # Try local data folder first, then fallback to job-application-mcp-server
        current_role_path = Path(__file__).parent / 'data' / 'current_role.json'
        if not current_role_path.exists():
            current_role_path = Path(__file__).parent.parent / 'job-application-mcp-server' / 'config' / 'current_role.json'
        if current_role_path.exists():
            with open(current_role_path) as f:
                current = json.load(f)
                exp = current.get('current_experience', {})
                if exp:
                    # Prepend current role to resume content
                    current_role_text = f"\n\nCURRENT ROLE at {exp.get('company')} ({exp.get('type', 'Contract')}):\n"
                    current_role_text += f"Location: {exp.get('location')}\nDates: {exp.get('dates')}\n\n"
                    for client in exp.get('clients', []):
                        title = client.get('title', 'Senior Software Engineer')
                        current_role_text += f"Client: {client['name']}\nTitle: {title}\nDates: {client['dates']}\n"
                        for point in client.get('points', []):
                            current_role_text += f"• {point}\n"
                        current_role_text += "\n"

                    # Add role selection guidance
                    role_selection = current.get('role_selection', {})
                    if role_selection:
                        current_role_text += "\nROLE SELECTION GUIDANCE:\n"
                        current_role_text += "- For Software Engineer positions: Show ONLY the Intuit client experience AND PayPal_SWE experience\n"
                        current_role_text += "- For QA/SDET/Test/Quality Engineer positions: Show ONLY the Apple client experience AND PayPal_QE experience\n"
                        current_role_text += "- Include only ONE version of each company based on the job title being applied for\n"
                        current_role_text += "- PayPal experience should show the role matching the job type (SWE or QE)\n"

                    resume_content = current_role_text + resume_content

                # Add missing experience (e.g., PayPal) - filter by role type
                # Role type will be injected by resume_api.py as __ROLE_TYPE__ marker
                missing_exp = current.get('missing_experience', {})
                if missing_exp:
                    missing_text = "\n\nMISSING EXPERIENCE (MUST include in resume - insert in correct chronological order):\n"
                    added_companies = set()

                    for name, details in missing_exp.items():
                        role_type = details.get('role_type', '')
                        company_base = details.get('company', name)

                        # Skip if we already added this company
                        if company_base in added_companies:
                            continue

                        # Include role_type marker so AI can select appropriately
                        if role_type:
                            missing_text += f"\n[USE FOR {role_type.upper()} ROLES]\n"
                        missing_text += f"{company_base}\n"
                        missing_text += f"Title: {details.get('title', '')}\n"
                        missing_text += f"Location: {details.get('location', '')}\n"
                        missing_text += f"Dates: {details.get('dates', '')}\n"
                        if details.get('technologies'):
                            missing_text += f"Technologies: {', '.join(details.get('technologies', []))}\n"
                        missing_text += "Achievements:\n"
                        for point in details.get('points', []):
                            missing_text += f"• {point}\n"

                        # Don't add to set - we want BOTH versions shown so AI can pick
                        # added_companies.add(company_base)

                    resume_content = missing_text + resume_content

                # Add experience restructure info (Altimetrik as parent company)
                restructure = current.get('experience_restructure', {})
                if restructure:
                    restructure_text = "\n\nEXPERIENCE RESTRUCTURE (show these as contract roles under parent company):\n"
                    for company, details in restructure.items():
                        restructure_text += f"\n{company} ({details.get('dates', '')}) - {details.get('type', '')}\n"
                        restructure_text += f"Clients: {', '.join(details.get('clients', []))}\n"
                        if 'coe' in details:
                            restructure_text += f"CoE Role: {details['coe'].get('role', '')}\n"
                            for point in details['coe'].get('points', []):
                                restructure_text += f"• {point}\n"
                    resume_content = restructure_text + resume_content

                # Add date corrections
                date_corrections = current.get('date_corrections', {})
                if date_corrections:
                    corrections_text = "\n\nDATE CORRECTIONS (use these dates instead):\n"
                    for company, dates in date_corrections.items():
                        corrections_text += f"- {company}: {dates}\n"
                    resume_content = corrections_text + resume_content

        return resume_content
    except Exception as e:
        print(f"Error loading resume: {e}")
        return None

def get_structured_resume(job_title: str, company: str, job_description: str, resume_content: str) -> dict:
    """Get AI-customized resume in structured JSON format"""
    
    prompt = """Analyze the resume and job description, then return a JSON object with this EXACT structure:
{
    "name": "Full Name",
    "title": "Professional Title tailored to job",
    "email": "mohammedmazhermd@gmail.com",
    "phone": "(510) 771-4493",
    "location": "San Jose, CA",
    "linkedin": "linkedin URL or empty string",
    "summary": "3-4 sentence professional summary tailored to the job",
    "skills": {
        "Languages": "Python, Java, JavaScript, etc",
        "Frameworks": "Spring Boot, React, etc",
        "Cloud & DevOps": "AWS, Kubernetes, Docker, etc",
        "Databases": "PostgreSQL, MongoDB, etc"
    },
    "experience": [
        {
            "title": "Job Title",
            "company": "Company Name",
            "location": "City, State",
            "dates": "MMM YYYY - MMM YYYY",
            "points": [
                "First bullet point achievement with metrics and numbers",
                "Second bullet point achievement with specific results",
                "Third bullet point showing technical skills used"
            ]
        }
    ]
}

CRITICAL - SKILLS FORMAT:
- Skills MUST be comma-separated strings like "Java, Python, JavaScript"
- Do NOT return skills as an array of characters
- Example correct format: "Languages": "Java, Python, JavaScript, TypeScript"

CRITICAL - EXPERIENCE POINTS:
- Each job MUST have a "points" array with 3-5 bullet points
- Points should be specific achievements from the original resume
- Include metrics, percentages, and specific technologies used
- Do NOT return empty points arrays

CRITICAL - CONTACT INFO (USE EXACTLY):
- Email: mohammedmazhermd@gmail.com (NOT muhammadkasim@gmail.com)
- Phone: (510) 771-4493
- Location: San Jose, CA

CRITICAL - EXPERIENCE ORDER:
- List experience in REVERSE CHRONOLOGICAL ORDER (most recent job FIRST)
- Galaxy I Tech (current) must be FIRST
- Zyme Solutions (oldest) must be LAST

NOTE: For contract/consulting roles with multiple clients under one parent company,
format as separate entries like:
- "Galaxy I Tech - Client: Apple"
- "Galaxy I Tech - Client: Intuit"
This shows the consulting relationship clearly while highlighting client work.

CRITICAL REQUIREMENTS:
- GALAXY I TECH CLIENT SELECTION (VERY IMPORTANT - SHOW ONLY ONE CLIENT):
  * If job title contains "QA", "SDET", "Test", "Quality": Show ONLY "Galaxy I Tech - Client: Apple"
  * If job title contains "Developer", "Engineer", "Architect" (without QA/Test): Show ONLY "Galaxy I Tech - Client: Intuit"
  * NEVER show both Apple and Intuit together! They are the same time period - pick ONE based on job type.

- Include ALL work experience from the original resume (every single job)
- ORDER EXPERIENCE STRICTLY BY DATE (most recent first). The correct order is:
  1. Galaxy I Tech (Jun 2025 - Present) - FIRST (current role) - show ONE client only!
  2. Citi (Oct 2024 - Jun 2025)
  3. LPL Financial (Oct 2023 - Sep 2024)
  4. Amazon (Sep 2022 - Oct 2023)
  5. Tusimple (Mar 2022 - Sep 2022)
  6. Proofpoint (Aug 2021 - Mar 2022)
  7. Altimetrik - PayPal (Mar 2020 - Aug 2021) - after Proofpoint!
  8. Altimetrik - Ancestry (Jan 2019 - Mar 2020)
  9. Altimetrik - Intuit (Jan 2014 - Jan 2019)
  10. Barclays (Nov 2012 - Dec 2013)
  11. Zyme Solutions (Jan 2007 - Nov 2012)
- DATE FORMAT IS MANDATORY: Always use "MMM YYYY - MMM YYYY" format (e.g., "Sep 2022 - Oct 2023")
  * Convert "2022-09-17" to "Sep 2022"
  * Convert "2023-10-31" to "Oct 2023"
  * Use "Present" for current roles
  * NEVER use YYYY-MM-DD format
- Include at least 3-4 bullet points per job
- Tailor bullet points to match the job description keywords
- Keep all information truthful from the original resume
- Prioritize skills mentioned in job description
- Quantify achievements where possible
- Return ONLY valid JSON, no other text"""

    # Get LLM configuration
    try:
        llm_config = get_llm_config()
    except ValueError as e:
        print(str(e))
        return None

    print(f"🤖 Using {llm_config['name']} ({llm_config['model']})")

    # Build request headers
    headers = {
        "Authorization": f"Bearer {llm_config['api_key']}",
        "Content-Type": "application/json"
    }

    # OpenRouter requires additional headers
    if llm_config['provider'] == 'openrouter':
        headers["HTTP-Referer"] = "https://github.com/resume-generator"
        headers["X-Title"] = "Resume Generator"

    # Build request payload
    payload = {
        "model": llm_config['model'],
        "messages": [
            {"role": "system", "content": prompt},
            {"role": "user", "content": f"Job: {job_title} at {company}\n\nJob Description:\n{job_description}\n\nResume:\n{resume_content}"}
        ],
        "max_tokens": 4000,
        "temperature": 0.5
    }

    # Add JSON response format (supported by Groq and Cerebras)
    if llm_config['provider'] in ['groq', 'cerebras']:
        payload["response_format"] = {"type": "json_object"}

    response = requests.post(llm_config['url'], headers=headers, json=payload)

    result = response.json()

    # Check for errors
    if "error" in result:
        error = result['error']
        if isinstance(error, dict):
            print(f"❌ API Error: {error.get('message', str(error))}")
        else:
            print(f"❌ API Error: {error}")
        return None

    if "choices" in result:
        try:
            content = result["choices"][0]["message"]["content"]
            print(f"📝 LLM Response (first 500 chars): {content[:500]}")
            data = json.loads(content)
            # Debug: print skills structure
            if 'skills' in data:
                print(f"🔧 Skills structure: {type(data['skills'])} - {data['skills']}")
            return data
        except json.JSONDecodeError as e:
            print(f"JSON parse error: {e}")
            print(f"Raw content: {result['choices'][0]['message']['content'][:500]}")
            return None
    return None

def generate_pdf(data: dict, output_path: str):
    """Generate PDF resume using ReportLab (industry standard)"""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    doc = SimpleDocTemplate(output_path, pagesize=letter,
                           leftMargin=0.6*inch, rightMargin=0.6*inch,
                           topMargin=0.5*inch, bottomMargin=0.5*inch)

    styles = getSampleStyleSheet()

    # Custom styles
    name_style = ParagraphStyle('Name', parent=styles['Title'], fontSize=22,
                                textColor=HexColor('#1e3a5f'), alignment=TA_CENTER, spaceAfter=4)
    title_style = ParagraphStyle('JobTitle', parent=styles['Normal'], fontSize=12,
                                 textColor=HexColor('#2563eb'), alignment=TA_CENTER, spaceAfter=4)
    contact_style = ParagraphStyle('Contact', parent=styles['Normal'], fontSize=9,
                                   textColor=HexColor('#6b7280'), alignment=TA_CENTER, spaceAfter=12)
    section_style = ParagraphStyle('Section', parent=styles['Heading2'], fontSize=11,
                                   textColor=HexColor('#1e3a5f'), spaceAfter=6, spaceBefore=10)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10,
                                textColor=HexColor('#374151'), spaceAfter=4, leading=14)
    exp_title_style = ParagraphStyle('ExpTitle', parent=styles['Normal'], fontSize=10,
                                     textColor=HexColor('#1a1a1a'), fontName='Helvetica-Bold')
    company_style = ParagraphStyle('Company', parent=styles['Normal'], fontSize=9,
                                   textColor=HexColor('#2563eb'), fontName='Helvetica-Oblique', spaceAfter=4)
    bullet_style = ParagraphStyle('Bullet', parent=styles['Normal'], fontSize=9,
                                  textColor=HexColor('#374151'), leftIndent=12, spaceAfter=2, leading=12)

    story = []

    # Header
    story.append(Paragraph(data.get('name', ''), name_style))
    story.append(Paragraph(data.get('title', ''), title_style))
    contact = f"{data.get('email', '')} &nbsp;|&nbsp; {data.get('phone', '')} &nbsp;|&nbsp; {data.get('location', '')}"
    story.append(Paragraph(contact, contact_style))

    # Summary
    story.append(Paragraph('PROFESSIONAL SUMMARY', section_style))
    story.append(Paragraph(data.get('summary', ''), body_style))

    # Skills
    story.append(Paragraph('TECHNICAL SKILLS', section_style))
    for category, skills in data.get('skills', {}).items():
        # Debug: print raw skills value
        print(f"📊 PDF Skills - {category}: type={type(skills)}, value={skills}")
        # Handle skills as list or string
        if isinstance(skills, list):
            skills_text = ', '.join(str(s) for s in skills)
        else:
            skills_text = skills  # Already a string, use directly
        print(f"📊 PDF Skills text: {skills_text}")
        story.append(Paragraph(f"<b>{category}:</b> {skills_text}", body_style))

    # Experience
    story.append(Paragraph('PROFESSIONAL EXPERIENCE', section_style))
    print(f"📋 Experience count: {len(data.get('experience', []))}")
    for job in data.get('experience', []):
        print(f"📋 Job: {job.get('title', '')} at {job.get('company', '')}")
        print(f"📋 Points: {job.get('points', [])}")
        story.append(Paragraph(f"<b>{job.get('title', '')}</b> | {job.get('dates', '')}", exp_title_style))
        story.append(Paragraph(f"{job.get('company', '')} - {job.get('location', '')}", company_style))
        for point in job.get('points', []):
            story.append(Paragraph(f"• {point}", bullet_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    print(f"✅ PDF saved: {output_path}")

def generate_docx(data: dict, output_path: str):
    """Generate DOCX resume"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Header
    name = doc.add_heading(data['name'], 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title_para = doc.add_paragraph(data['title'])
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph(f"{data['email']} | {data['phone']} | {data['location']}")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(data['summary'])
    
    # Skills
    doc.add_heading('Technical Skills', level=1)
    for category, skills in data['skills'].items():
        # Handle skills as list or string
        if isinstance(skills, list):
            # Check if it's a list of single characters (LLM bug) - join without separator
            if skills and all(len(str(s)) <= 2 for s in skills):
                skills_text = ''.join(str(s) for s in skills)
            else:
                skills_text = ', '.join(str(s) for s in skills)
        else:
            skills_text = str(skills)
        p = doc.add_paragraph()
        p.add_run(f"{category}: ").bold = True
        p.add_run(skills_text)
    
    # Experience
    doc.add_heading('Professional Experience', level=1)
    for job in data['experience']:
        p = doc.add_paragraph()
        p.add_run(f"{job['title']}").bold = True
        p.add_run(f" | {job['company']} | {job['dates']}")
        for point in job['points']:
            doc.add_paragraph(point, style='List Bullet')
    
    doc.save(output_path)
    print(f"✅ DOCX saved: {output_path}")

def main():
    print("\n" + "="*60)
    print("📄 PROFESSIONAL RESUME GENERATOR")
    print("   Multi-LLM Support: Groq | Cerebras | OpenRouter")
    print("="*60 + "\n")

    # Show current LLM provider
    provider = os.environ.get("LLM_PROVIDER", "groq").lower()
    if provider in LLM_PROVIDERS:
        config = LLM_PROVIDERS[provider]
        print(f"🤖 LLM Provider: {config['name']} ({config['model']})")
        print(f"   Switch with: export LLM_PROVIDER=cerebras|groq|openrouter\n")

    resume = load_resume()
    if not resume:
        print("❌ Could not load resume")
        return
    print(f"✅ Resume loaded\n")
    
    job_title = input("Job Title: ").strip() or "Senior Software Engineer"
    company = input("Company: ").strip() or "Google"
    print("Job Description (paste, then press Enter twice):")
    lines = []
    while True:
        line = input()
        if line == "" and lines and lines[-1] == "":
            break
        lines.append(line)
    job_description = "\n".join(lines[:-1]) if lines else "Software Engineer role"
    
    print("\n⏳ AI is customizing your resume...\n")
    data = get_structured_resume(job_title, company, job_description, resume[:8000])
    
    if not data:
        print("❌ Failed to generate structured resume")
        return
    
    output_dir = Path(__file__).parent / 'output'
    output_dir.mkdir(exist_ok=True)
    
    safe_company = company.lower().replace(' ', '_').replace('.', '')
    generate_pdf(data, str(output_dir / f"resume_{safe_company}.pdf"))
    generate_docx(data, str(output_dir / f"resume_{safe_company}.docx"))
    
    print(f"\n🎉 Done! Check the 'output' folder.")

if __name__ == "__main__":
    main()

